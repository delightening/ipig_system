"""抽取 .docx 完整格式資訊：字體、字級、間距、頁邊距、表格。"""
import sys
from docx import Document

EMU_PER_PT = 12700
EMU_PER_CM = 360000


def emu_to_cm(emu):
    return round(emu / EMU_PER_CM, 2) if emu is not None else None


def emu_to_pt(emu):
    return round(emu / EMU_PER_PT, 1) if emu is not None else None


def fmt_para(p, idx):
    pf = p.paragraph_format
    runs_info = []
    for r in p.runs:
        if not r.text.strip():
            continue
        font = r.font
        sz = font.size.pt if font.size is not None else None
        runs_info.append({
            "text": r.text[:40] + ("..." if len(r.text) > 40 else ""),
            "font": font.name,
            "size_pt": sz,
            "bold": font.bold,
            "italic": font.italic,
        })
    return {
        "idx": idx,
        "style": p.style.name if p.style is not None else None,
        "text_preview": p.text[:60].replace("\n", " "),
        "alignment": str(pf.alignment).split(".")[-1] if pf.alignment is not None else None,
        "line_spacing": pf.line_spacing,
        "space_before_pt": pf.space_before.pt if pf.space_before is not None else None,
        "space_after_pt": pf.space_after.pt if pf.space_after is not None else None,
        "first_line_indent_pt": pf.first_line_indent.pt if pf.first_line_indent is not None else None,
        "left_indent_pt": pf.left_indent.pt if pf.left_indent is not None else None,
        "runs": runs_info,
    }


def main(path):
    doc = Document(path)

    # Section / page setup
    print("=" * 70)
    print("SECTIONS (頁面設定)")
    print("=" * 70)
    for i, sec in enumerate(doc.sections):
        print(f"\n[Section {i}]")
        print(f"  page_height:  {emu_to_cm(sec.page_height)} cm")
        print(f"  page_width:   {emu_to_cm(sec.page_width)} cm")
        print(f"  orientation:  {sec.orientation}")
        print(f"  top_margin:   {emu_to_cm(sec.top_margin)} cm")
        print(f"  bottom_margin:{emu_to_cm(sec.bottom_margin)} cm")
        print(f"  left_margin:  {emu_to_cm(sec.left_margin)} cm")
        print(f"  right_margin: {emu_to_cm(sec.right_margin)} cm")
        print(f"  header_dist:  {emu_to_cm(sec.header_distance)} cm")
        print(f"  footer_dist:  {emu_to_cm(sec.footer_distance)} cm")
        print(f"  gutter:       {emu_to_cm(sec.gutter)} cm")

    # Styles
    print("\n" + "=" * 70)
    print("STYLES (樣式定義)")
    print("=" * 70)
    for style in doc.styles:
        if not hasattr(style, "font"):
            continue
        font = style.font
        if font.name is None and font.size is None:
            continue
        sz = font.size.pt if font.size is not None else None
        print(f"  [{style.type}] {style.name:30s} font={font.name!s:20s} size={sz}pt bold={font.bold}")

    # Paragraphs
    print("\n" + "=" * 70)
    print(f"PARAGRAPHS (共 {len(doc.paragraphs)} 段)")
    print("=" * 70)
    for i, p in enumerate(doc.paragraphs):
        info = fmt_para(p, i)
        if not info["text_preview"].strip() and not info["runs"]:
            continue
        print(f"\n[{i}] style={info['style']}  align={info['alignment']}")
        print(f"    TEXT: {info['text_preview']}")
        spacing_bits = []
        if info["space_before_pt"] is not None:
            spacing_bits.append(f"before={info['space_before_pt']}pt")
        if info["space_after_pt"] is not None:
            spacing_bits.append(f"after={info['space_after_pt']}pt")
        if info["line_spacing"] is not None:
            spacing_bits.append(f"line={info['line_spacing']}")
        if info["first_line_indent_pt"] is not None:
            spacing_bits.append(f"indent={info['first_line_indent_pt']}pt")
        if info["left_indent_pt"] is not None:
            spacing_bits.append(f"left={info['left_indent_pt']}pt")
        if spacing_bits:
            print(f"    SPACING: {', '.join(spacing_bits)}")
        for r in info["runs"]:
            font_bits = []
            if r["font"]:
                font_bits.append(f"font={r['font']}")
            if r["size_pt"] is not None:
                font_bits.append(f"size={r['size_pt']}pt")
            if r["bold"]:
                font_bits.append("bold")
            if r["italic"]:
                font_bits.append("italic")
            if font_bits:
                print(f"    RUN: [{' '.join(font_bits)}] {r['text']}")

    # Tables
    print("\n" + "=" * 70)
    print(f"TABLES (共 {len(doc.tables)} 個)")
    print("=" * 70)
    for ti, table in enumerate(doc.tables):
        style_name = table.style.name if table.style is not None else None
        print(f"\n[Table {ti}]  rows={len(table.rows)}  cols={len(table.columns)}  style={style_name}")
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                text = cell.text.replace("\n", " | ")[:50]
                if text.strip():
                    print(f"  [{ri},{ci}] {text}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(f"Usage: python {sys.argv[0]} <path_to_docx>", file=sys.stderr)
        sys.exit(1)
    main(sys.argv[1])
